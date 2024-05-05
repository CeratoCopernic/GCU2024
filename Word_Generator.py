from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor
import os

def generate_word(init_path, output_folder):
    with open(init_path, "r") as file:
        contenu = file.readlines()

    # Créer un nouveau document Word
    doc = Document()

    # Fonction pour ajouter un paragraphe centré
    def ajouter_paragraphe_centre(texte, style=None):
        paragraphe = doc.add_paragraph(style=style)
        paragraphe.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraphe.paragraph_format.space_after = Pt(0)  # Supprimer l'espacement après
        paragraphe.paragraph_format.space_before = Pt(0)
        paragraphe.add_run(texte)

    def ajouter_paragraphe_norm(texte, style=None):
        paragraphe = doc.add_paragraph(style=style)
        #paragraphe.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraphe.paragraph_format.space_after = Pt(0)  # Supprimer l'espacement après
        paragraphe.paragraph_format.space_before = Pt(0)
        paragraphe.add_run(texte)

    # Fonction pour ajouter un titre
    def ajouter_titre(texte):
        paragraphe = doc.add_paragraph(style="Heading 1")
        paragraphe.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraphe.paragraph_format.space_after = Pt(0)  # Supprimer l'espacement après
        paragraphe.add_run(texte)

    # Fonction pour ajouter un sous-titre
    def ajouter_sous_titre(texte):
        paragraphe = doc.add_paragraph(style="Heading 2")
        paragraphe.paragraph_format.space_after = Pt(0)  # Supprimer l'espacement après
        paragraphe.add_run(texte)

    def mettre_en_gras(texte):
        paragraphe = doc.add_paragraph()
        #paragraphe.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraphe.paragraph_format.space_after = Pt(0)  # Supprimer l'espacement après
        run = paragraphe.add_run(texte)
        run.bold = True
    
    def mettre_en_vert(texte):
        paragraphe = doc.add_paragraph()
        #paragraphe.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraphe.paragraph_format.space_after = Pt(0)  # Supprimer l'espacement après
        run = paragraphe.add_run(texte)
        run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)

    def mettre_en_rouge(texte):
        paragraphe = doc.add_paragraph()
        paragraphe.paragraph_format.space_after = Pt(0)  # Supprimer l'espacement après
        run = paragraphe.add_run(texte)
        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # Couleur rouge


    # Ajouter le contenu du fichier texte au document Word
    for ligne in contenu:
        if ligne.startswith("JEU PERMANENT"):
            ajouter_paragraphe_centre(ligne.strip(), "Title")
        elif ligne.startswith("FICHE RÉCAPITULATIVE"):
            ajouter_titre(ligne.strip())
        elif ligne.startswith("Mise à jour"):
            ajouter_paragraphe_centre(ligne.strip())
        elif ligne.startswith(("1.","2.","3.")):
            ajouter_sous_titre(ligne.strip())
        elif ligne.startswith(("Territoires possédés :","Ressources :","Armée :","État du mar")):
            mettre_en_gras(ligne.strip())
        elif ligne.strip().isdigit():
            ajouter_sous_titre(ligne.strip())
        elif ligne.startswith("Territoire n°") or ligne.startswith("----"):
            mettre_en_gras(ligne.strip())
        elif ligne.startswith(("SERV","HAR","ARCH","LANC","AUTOC","VAILL","PAL","TEMP","GUE","PRE","CHE")):
            mettre_en_gras(ligne.strip())
        elif ligne.strip().endswith("Oui"):
            mettre_en_vert(ligne.strip())
        else:
            ajouter_paragraphe_norm(ligne.strip())

    # Nom du fichier de sortie
    nom_fichier = os.path.splitext(os.path.basename(init_path))[0]  # Nom du fichier sans extension
    nom_fichier += ".docx"  # Ajouter l'extension .docx

    # Chemin complet du fichier de sortie
    chemin_sortie = os.path.join(output_folder, nom_fichier)

    # Enregistrer le document Word dans le dossier de sortie
    doc.save(chemin_sortie)

dossier_entree = "/Users/alexandredemerode/Desktop/Jeu Perm - GCU 2024/Messages"
dossier_sortie = "/Users/alexandredemerode/Desktop/Jeu Perm - GCU 2024/Words"

# Créer le dossier de sortie s'il n'existe pas
if not os.path.exists(dossier_sortie):
    os.makedirs(dossier_sortie)

# Parcourir tous les fichiers .txt dans le dossier d'entrée et générer les fichiers Word correspondants dans le dossier de sortie
for fichier in os.listdir(dossier_entree):
    if fichier.endswith(".txt"):
        chemin_fichier = os.path.join(dossier_entree, fichier)
        generate_word(chemin_fichier, dossier_sortie)

in_pth = "/Users/alexandredemerode/Desktop/Jeu Perm - GCU 2024/Catastrophes/General_Report.txt"
out_doss = "/Users/alexandredemerode/Desktop/Jeu Perm - GCU 2024/Catastrophes"
generate_word(in_pth,out_doss)
